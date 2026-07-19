from pathlib import Path
print("IMPORT 1")
import io

print("IMPORT 2")


print("IMPORT 3")
import pdfplumber

print("IMPORT 4")
from docx import Document

print("IMPORT 5")
import PyPDF2

print("IMPORT 6")
from backend.utils.file_utils import (
    FileParsingError,
    TextExtractionError,
    FileUploadError,
    log_error,
    log_warning,
    log_info,
    with_fallback,
)

print("IMPORT 7")
from backend.core.config import (
    MAX_FILE_SIZE_BYTES,
    MAX_FILE_SIZE_MB,
    SUPPORTED_MIME_TYPES,
)

print("IMPORT FINISHED")

class FileValidationError(Exception):
    pass

from typing import Tuple, Optional

from typing import Tuple, Optional
from pathlib import Path

def validate_file(file_data: bytes, filename: str) -> Tuple[bool, str, Optional[str]]:
    print("VALIDATE 1")

    file_size_bytes = len(file_data)
    print("VALIDATE 2")

    if file_size_bytes > MAX_FILE_SIZE_BYTES:
        size_mb = file_size_bytes / (1024 * 1024)
        return False, (
            f"File size ({size_mb:.2f} MB) exceeds the maximum of {MAX_FILE_SIZE_MB} MB."
        ), None

    if file_size_bytes == 0:
        return False, "Uploaded file is empty.", None

    print("VALIDATE 3")

    extension = Path(filename).suffix.lower()

    extension_map = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "doc",
    }

    if extension not in extension_map:
        return False, f"Unsupported file type: {extension}", None

    print("VALIDATE 4")

    return True, "", extension_map[extension]

def _extract_pdf_hyperlinks(file_data: bytes) -> str:
    urls = []
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_data))
        for page in reader.pages:
            if '/Annots' not in page:
                continue
            for annot_ref in page['/Annots']:
                try:
                    annot = annot_ref.get_object()
                    if annot.get('/Subtype') != '/Link':
                        continue
                    action = annot.get('/A', {})
                    uri = action.get('/URI', '')
                    if uri and isinstance(uri, (str, bytes)):
                        # PyPDF2 may return bytes for URI values
                        if isinstance(uri, bytes):
                            uri = uri.decode('utf-8', errors='ignore')
                        uri = uri.strip()
                        if uri.startswith('http'):
                            urls.append(uri)
                except Exception:
                    pass
    except Exception:
        pass
    return '\n'.join(urls)


def _extract_pdf_with_pdfplumber(file_data: bytes) -> str:
    print("PDFPLUMBER 1")

    text = ""

    print("PDFPLUMBER 2 - Opening PDF")

    with pdfplumber.open(io.BytesIO(file_data)) as pdf:

        print("PDFPLUMBER 3 - PDF opened")

        print("Total Pages:", len(pdf.pages))

        for i, page in enumerate(pdf.pages):

            print(f"Reading page {i+1}")

            print("Calling page.extract_text()")

            page_text = page.extract_text()

            print("Returned from extract_text()")

            if page_text:
                text += page_text + "\n"

            print(f"Finished page {i+1}")

    print("PDFPLUMBER 4")

    if not text.strip():
        raise TextExtractionError("No text extracted")

    print("PDFPLUMBER 5")

    return text

def _extract_pdf_with_pypdf2(file_data: bytes) -> str:
    print("PYPDF2 1")

    text = ""

    reader = PyPDF2.PdfReader(io.BytesIO(file_data))

    print("PYPDF2 2")

    for i, page in enumerate(reader.pages):

        print("PYPDF2 Page", i + 1)

        page_text = page.extract_text()

        print("PYPDF2 Done", i + 1)

        if page_text:
            text += page_text + "\n"

    print("PYPDF2 3")

    return text


def extract_text_from_pdf(file_data: bytes) -> str:
    print("PDF 1 - Enter extract_text_from_pdf")

    try:
        print("PDF 2 - Before with_fallback")

        result, used_fallback = with_fallback(
            _extract_pdf_with_pdfplumber,
            _extract_pdf_with_pypdf2,
            file_data,
            log_fallback=True,
        )

        print("PDF 3 - After with_fallback")

        if used_fallback:
            print("PDF FALLBACK USED")

        return result

    except Exception as e:
        print("PDF ERROR:", repr(e))
        raise
    

def extract_text_from_docx(file_data: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_data))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        text = '\n'.join(text_parts)

        if not text.strip():
            raise FileParsingError(
                'No text could be extracted from the document. '
                'The document may be empty or corrupted.'
            )
        
        try:
            for rel in doc.part.rels.values():
                if 'hyperlink' in rel.reltype.lower():
                    url = rel._target
                    if isinstance(url, str) and url.startswith('http'):
                        text += '\n' + url
        except Exception:
            pass

        log_info(f'Extracted {len(text)} chars from DOCX', context='resume_parser')
        return text.strip()

    except FileParsingError:
        raise   # Re-raise unchanged — don't wrap in another FileParsingError

    except Exception as e:
        log_error(e, context='extract_text_from_docx')
        raise FileParsingError(
            'Failed to extract text from DOCX. '
            'The document may be corrupted or in an unsupported format. '
            'Please try re-saving or converting to PDF.'
        ) from e

def extract_text_from_doc(file_data: bytes) -> str:
    raise FileParsingError(
        'Legacy .doc format is not supported. '
        'Please convert your document to .docx or .pdf and try again. '
        'You can convert using Microsoft Word, Google Docs, or online tools.'
    )

def extract_text(file_data: bytes, file_type: str) -> str:

    print("EXTRACT 1:", file_type)

    if file_type == "pdf":

        print("EXTRACT 2 - PDF")

        result = extract_text_from_pdf(file_data)

        print("EXTRACT 3")

        return result

    elif file_type == "docx":

        print("EXTRACT DOCX")

        return extract_text_from_docx(file_data)

    elif file_type == "doc":

        print("EXTRACT DOC")

        return extract_text_from_doc(file_data)

    raise FileValidationError(
        f"Invalid file type: {file_type}"
    )
    
def parse_resume_file(file_data: bytes, filename: str):

    print("PARSER 1")

    is_valid, error_msg, file_type = validate_file(file_data, filename)

    print("PARSER 2")

    if not is_valid:
        raise FileValidationError(error_msg)

    print("PARSER 3")

    print("Calling extract_text()")

    text = extract_text(file_data, file_type)

    print("Returned from extract_text()")

    print("PARSER 4")

    metadata = {
        "filename": filename,
        "file_type": file_type,
        "text_length": len(text),
    }

    print("PARSER 5")

    return text, metadata